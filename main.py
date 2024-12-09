import os
import numpy as np
import pandas as pd
import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, Trainer, TrainingArguments
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from torch.utils.data import Dataset

class WritingStyleRAG:
    def __init__(self, user_essays_dir):
        """
        Initialize the Writing Style RAG system
        
        Args:
            user_essays_dir (str): Directory containing user's original essays
        """
        self.user_essays_dir = user_essays_dir

        # Preprocessing configurations
        self.max_length = 512  # Maximum token length

        # Embedding model for style capture
        self.style_embedding_model = SentenceTransformer('all-MiniLM-L6-v2', local_files_only=True)

        # Language generation model
        self.tokenizer = AutoTokenizer.from_pretrained('facebook/bart-large', local_files_only=True)
        self.generation_model = AutoModelForSeq2SeqLM.from_pretrained('facebook/bart-large', local_files_only=True)

        # Style-related attributes
        self.style_embeddings = []
        self.essay_texts = []

    def load_essays(self):
        """
        Load essays from the specified directory
        """
        self.essay_texts = []
        self.style_embeddings = []

        for filename in os.listdir(self.user_essays_dir):
            # Skip temporary Word files
            if filename.startswith('~$'):
                continue

            file_path = os.path.join(self.user_essays_dir, filename)
            try:
                if filename.endswith('.docx'):
                    doc = Document(file_path)
                    essay_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                elif filename.endswith(('.txt', '.md')):
                    with open(file_path, 'r', encoding='utf-8') as file:
                        essay_text = file.read()

                # Process the text regardless of file type
                self.essay_texts.append(essay_text)
                style_embedding = self.style_embedding_model.encode(essay_text)
                self.style_embeddings.append(style_embedding)
                print(f"Successfully loaded: {filename}")
            except Exception as e:
                print(f"Error loading {filename}: {str(e)}")
                continue

        return len(self.essay_texts)

    def calculate_style_similarity(self, source_text):
        """
        Calculate similarity between source text and stored essay styles
        
        Args:
            source_text (str): Text to match against stored styles
        
        Returns:
            dict: Similarity scores for each stored essay
        """
        source_embedding = self.style_embedding_model.encode(source_text)

        similarities = []
        for stored_embedding in self.style_embeddings:
            sim = cosine_similarity([source_embedding], [stored_embedding])[0][0]
            similarities.append(sim)

        return {
            'similarities': similarities,
            'most_similar_index': np.argmax(similarities)
        }

    def style_transfer(self, source_text, top_k=1):
        """
        Transfer writing style to the source text
        
        Args:
            source_text (str): Text to be stylized
            top_k (int): Number of most similar styles to consider
        
        Returns:
            dict: Contains stylized text and similarity score
        """
        # Find most similar existing essays
        style_match = self.calculate_style_similarity(source_text)
        reference_text = self.essay_texts[style_match['most_similar_index']]

        # Create explicit style transfer prompt
        prompt = f"""
        Style: {reference_text[:200]}...
        
        Rewrite the following text in the same style as shown above:
        {source_text}
        
        Styled version:
        """

        # Tokenize with explicit prompt
        inputs = self.tokenizer(
            prompt,
            return_tensors='pt',
            max_length=self.max_length,
            truncation=True
        )

        # Generate with more controlled parameters
        outputs = self.generation_model.generate(
            input_ids=inputs['input_ids'],
            attention_mask=inputs['attention_mask'],
            max_length=self.max_length,
            num_return_sequences=1,
            temperature=0.7,
            top_p=0.9,           # Nucleus sampling
            repetition_penalty=1.2,  # Reduce repetition
            length_penalty=1.0,   # Encourage complete sentences
            do_sample=True       # Enable sampling for more natural text
        )

        # Decode and clean up the output
        stylized_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
        # Remove the prompt if it appears in the output
        stylized_text = stylized_text.replace(prompt.strip(), "").strip()

        return {
            'stylized_text': stylized_text,
            'style_similarity': style_match['similarities'][style_match['most_similar_index']],
            'reference_style': reference_text[:200]  # Return sample of reference style
        }

def read_source_text(filename):
    """Read source text from a file (.txt, .md, or .docx)"""
    if not filename.endswith(('.txt', '.md', '.docx')):
        raise ValueError("Source file must be .txt, .md, or .docx format")

    if filename.endswith('.docx'):
        doc = Document(filename)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        with open(filename, 'r', encoding='utf-8') as file:
            return file.read()

def save_stylized_text(result, output_filename='stylized_output.docx'):
    """Save stylized text to a Word document"""
    doc = Document()
    doc.add_heading('Stylized Text', 0)
    doc.add_paragraph(result['stylized_text'])
    
    # Add style metrics
    doc.add_heading('Style Metrics', level=1)
    doc.add_paragraph(f"Style Similarity Score: {result['style_similarity']:.2f}")
    
    doc.save(output_filename)
def main():
    rag_styler = WritingStyleRAG('training_essays')
    
    # Load user's essays
    num_essays = rag_styler.load_essays()
    print(f"Loaded {num_essays} essays for style learning")

    # Read source text from file (now supports multiple formats)
    source_text = read_source_text('source.docx')  # Can also use source.md or source.docx
    source_text = read_source_text('source.docx')
    result = rag_styler.style_transfer(source_text)

    print("Stylized Text:", result['stylized_text'])
    print("Style Similarity Score:", result['style_similarity'])
    # Save to Word document
    save_stylized_text(result)
    print(f"Stylized text saved to 'stylized_output.docx'")

if __name__ == '__main__':
    main()

# Additional Considerations:
# 1. Fine-tuning the model on user's specific writing style
# 2. Handling different essay types and genres
# 3. Implementing more sophisticated style embedding techniques